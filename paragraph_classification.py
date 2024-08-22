input("프로그램을 시작하려면 아무 키나 누르세요.")

tests = [
    "Conducted Emission - Voltage Method",
    "Conducted Emission - Current Method",
    "Radiated Emission",
    "Magnetic Field Emission",
    "CV test",
    "Transient Emission",
    "Bulk Current Injection: Immunity",
    "ALSE: Immunity",
    "ALSE: Radar Pulse",
    "Portable Transmitter Test",
    "Magnetic Field Immunity",
    "Transient Immunity: Supply lines",
    "Transient Immunity; Signal lines",
    "Electrostatic Discharge (ESD) - Unpowered Discharge",
    "Electrostatic Discharge (ESD) - Powered-up Discharge",
    "Test to be performed as shwon"
]


def read_test_info_from_file(file_path):
    with open(file_path, 'r') as file:
        lines = file.readlines()

    # Parse each line, splitting by comma and stripping whitespace
    additional_test_info = [line.strip().split(', ') for line in lines]

    # Convert empty strings to empty lists
    additional_test_info = [info if info != [''] else [] for info in additional_test_info]

    return additional_test_info


def read_config(config_path):
    config = {}
    sections = []
    section_mode = False
    with open(config_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
        for line in lines:
            line = line.strip()
            if line.startswith('file_type'):
                config['file_type'] = line.split(': ')[1]
            elif line.startswith('file_path'):
                config['file_path'] = line.split(': ')[1]
            elif line.startswith('start_page'):
                config['start_page'] = int(line.split(': ')[1])
            elif line.startswith('end_page'):
                config['end_page'] = int(line.split(': ')[1])
            elif line.startswith('sections'):
                # sections 항목 이후의 줄들을 저장
                section_mode = True
            elif section_mode:
                sections.append(line)
    
    config['sections'] = sections
    return config
import fitz 
import copy
import re
from docx import Document
import sys
from docx.oxml.ns import qn
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from rake_nltk import Rake
import nltk
config_file_path = "config.txt"
config = read_config(config_file_path)

file_type = config['file_type']
file_path = config['file_path']
start_page = config['start_page']
end_page = config['end_page']
sections = config['sections']

print(f"file_type: {file_type}")
print(f"file_path: {file_path}")
print(f"start_page: {start_page}")
print(f"end_page: {end_page}")
print(f"sections: {sections}")

def read_test_info_from_file(file_path):
    with open(file_path, 'r') as file:
        lines = file.readlines()

    # Parse each line, splitting by comma and stripping whitespace
    additional_test_info = [line.strip().split(', ') for line in lines]

    # Convert empty strings to empty lists
    additional_test_info = [info if info != [''] else [] for info in additional_test_info]

    return additional_test_info

# Usage


parameter_file_path = 'selected_parameters.txt'
additional_test_info = read_test_info_from_file(parameter_file_path)
additional_test_info.append("")


tests_input = copy.deepcopy(tests)
for i in range(len(tests_input)):
    for j in additional_test_info[i]:
        tests_input[i] += " " + j


def remove_hanja(text):
    hanja_pattern = re.compile(r'[\u4E00-\u9FFF\u3400-\u4DBF]')
    cleaned_text = hanja_pattern.sub('', text)
    return cleaned_text


def split_number_string(s):
    #match = re.match(r"((?:\d+\.)+)(\S+)", s)'
    match = re.match(r"((?:\d+\.(?!\d))+)\s*(\S+)", s)
    if match:
        number_part = match.group(1)
        string_part = match.group(2)
        result = number_part + " "+string_part 
        result = re.sub(r'。', ' ', result)
        result = re.sub(r'\s{2,}', ' ', result)
        return result
    else:
        return ""


def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pages.append(page.get_text())
    new_pages = []
    for page in pages:
    
        page = remove_hanja(page)
        page = page.replace("\t", " ").replace("\n", " ").strip().lower()
        #page = page.replace("-", " ")
        page = re.sub(r'。', ' ', page)
        page = re.sub(r'\s{2,}', ' ', page).strip()

        new_pages.append(page)

    return new_pages

def split_paragraph_pages(pages, section):
    result = []
    current_section = 0
    txt = ""   
    past_words = ""
    for page in pages:
        while current_section <= len(section) - 1:
            preprocess_section = remove_hanja(section[current_section])
            preprocess_section = preprocess_section.replace("\t", " ").replace("\n"," ").strip().lower()
            #preprocess_section = preprocess_section.replace("-", " ")
            preprocess_section = preprocess_section.replace('。', " ")
            section_list = re.sub(r'\s{2,}', ' ', preprocess_section).strip().split(' ')
            section_list = [section_list]
                        
            section2 = split_number_string(section[current_section])
            if (len(section2)  != 0):
                preprocess_section2 = remove_hanja(section2)
                preprocess_section2 = preprocess_section2.replace("\t", " ").replace("\n"," ").strip().lower()
                preprocess_section2 = preprocess_section2.replace("-", " ")
                section_list2 = re.sub(r'\s{2,}', ' ', preprocess_section2).strip().split(' ')
                section_list.append(section_list2)
            # Find the position of the section in the current page
            words = page.lower().split()
            start_idx = 0
            found = False
            break_point = True

            for s in section_list:
                for i in range(len(words)):
                    #print(words[i:i+len(s)], "와 ", s, "비교중")
                    if words[i:i+len(s)] == s:
                        found = True
                        start_idx = i
                        break

            if found:
                # If section is found in the middle of the page
                part1 = ' '.join(words[:start_idx])
                part2 = ' '.join(words[start_idx:])

                if (part1.strip() or txt.strip()):
                    result.append(txt.strip()+part1.strip())
                
                # Set the remaining part of the page to be processed
                page = part2
                txt = ""
                current_section += 1
                
            else:
                break
        
        txt += page + "\n"

    if txt.strip():
        result.append(txt.strip())

    return result

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    new_pages = []

    for element in doc.element.body:
        if element.tag == qn('w:p'):  # Paragraph
            paragraph = element
            text = paragraph.text
            text = remove_hanja(text)
            text = text.replace("\t", " ").replace("\n", " ").strip().lower()
            text = re.sub(r'\s{2,}', ' ', text).strip()
            new_pages.append(text)
        elif element.tag == qn('w:tbl'):  # Table
            table_text = ""
            table = element
            for row in table.findall('.//w:tr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                row_text = ""
                for cell in row.findall('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    cell_text = ''.join(node.text for node in cell if node.text)
                    row_text += cell_text + '\t'
                row_text = row_text.strip() + '\n'
                table_text += row_text
            table_text += '\n'
            new_pages.append(table_text)
    return new_pages

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    text = '\n'.join(full_text)

    result = text.split("\n")
    with open("total_text.txt", 'w', encoding='utf-8') as f:
        f.write(text)

    with open('temp_result.txt', 'w', encoding='utf-8') as outfile:
        outfile.write(text)
    return result
print("file_path: ", file_path)
if(file_type == 'pdf'):
    pages = extract_text_from_pdf(file_path)
elif(file_type == 'docx'):
    pages = extract_text_from_docx(file_path)
    start_page = 2
    end_page = len(pages)
else :
    print("요청한 파일 형식이 잘못 되었습니다.")
    sys.exit(0)

if(file_type == 'pdf'):
    result = split_paragraph_pages(pages[start_page-1:end_page], sections)
elif(file_type == 'docx'):
    result = split_paragraph_pages(pages, sections)
result.pop(0)

if(len(sections) == len(result)):
    print("주어진 section을 모두 찾았습니다.")
else :
    print(f"{sections[len(result)]}부터 찾지 못했습니다.")
for i in result:
    print(i)
    print("------------")


stop_words = nltk.corpus.stopwords.words('english')
add_stopwords = ['test', 'testing','requirements', 'shall']
for i in add_stopwords:
    stop_words.append(i)
# 키워드 추출 (RAKE)
rake = Rake(stopwords = stop_words)
section_keywords = []
for part in result:
    rake.extract_keywords_from_text(part)
    keywords = rake.get_ranked_phrases()
    section_keywords.append(' '.join(keywords))

# TF-IDF 벡터화
vectorizer = TfidfVectorizer()
tfidf_matrix = vectorizer.fit_transform(section_keywords + tests_input)

# 유사도 계산
cosine_sim = cosine_similarity(tfidf_matrix[:len(result)], tfidf_matrix[len(result):])
print(cosine_sim)
# 각 섹션을 가장 유사한 테스트에 매칭
matched_tests = []

# for idx, similarities in enumerate(cosine_sim):
#     for i in range(len(test_keyword)):
#         total = 0 
#         for j in range(len(test_keyword[i])):
#              total += count_matches(sections[idx], test_keyword[i][j])

tests_match = {} # key : test, value : section의 목록
for i in tests:
    tests_match[i] = []
for idx, similarities in enumerate(cosine_sim):
    if idx == len(result):
        break
    most_similar_test_idx = similarities.argmax()
    matched_tests.append((result[idx], tests[most_similar_test_idx]))
    tests_match[tests[most_similar_test_idx]].append(sections[idx])

# 매칭 결과 출력
result_txt = ""

from collections import defaultdict
result_dict = {} # key : test, value : section의 내용들
for i in tests:
    result_dict[i] = []
i = 0
for section, test in matched_tests:
    print(f"{sections[i]} : {test}\n")
    result_dict[test].append(result[i])
    result_txt+= f"{sections[i]} : {test}\n"
    i += 1
    if(i == (len(sections) or len(result))):
        break
f = open("result_classification.txt", 'w', encoding='utf-8')
f.write(result_txt)

import os
import shutil

def create_folder(path):
    if os.path.exists(path):
        shutil.rmtree(path)
    os.makedirs(path)

folder_path = "classification_contents"
create_folder(folder_path)

for i in range(len(tests)):
    test = tests[i]
    for j in range(len(tests_match[test])):
        cleaned_test = re.sub(r'[^A-Za-z0-9\s]', '', test)
        new_folder_path = folder_path + '/' + cleaned_test
        create_folder(new_folder_path)
    for j in range(len(tests_match[test])):
        cleaned_test = re.sub(r'[^A-Za-z0-9\s]', '', test)
        new_folder_path = folder_path + '/' + cleaned_test
        section = tests_match[test][j]
        with open(f"{new_folder_path}/{j}.txt", 'w', encoding='utf-8') as f:
            f.write(section + "\n\n" + result_dict[test][j])


