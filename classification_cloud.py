from bayes_opt import BayesianOptimization
import numpy as np
import copy
import os

import fitz 
import copy
import re
from docx import Document
import sys 
from docx.oxml.ns import qn
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from rake_nltk import Rake
import nltk
from nltk.corpus import stopwords
from collections import Counter
import string


file_path_list = []
file_type_list = []
start_page_list = []
end_page_list = []
sections_list = []
answer_list = []


def read_optimizaton_data(folder_path):
    for filename in os.listdir(folder_path):
        if(filename.endswith(".txt")):
            file_path = os.path.join(folder_path, filename)
            print(f"Reading data from {file_path}")
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()

                file_path_item = lines[0].strip()
                file_type_item = lines[1].strip()
                start_page_item = int(lines[2].strip())
                end_page_item = int(lines[3].strip())
                answers = []
                sections = []
                for i in range(4, len(lines), 2):
                    section_item = lines[i].strip()
                    if i+1<len(lines):    
                        answer_item = [int(x) for x in lines[i+1].strip().split()]
                    else:
                        answer_item = []
                    sections.append(section_item)
                    answers.append(answer_item)
                sections_list.append(sections)
                answer_list.append(answers)
                file_path_list.append(file_path_item)
                file_type_list.append(file_type_item)
                start_page_list.append(start_page_item)
                end_page_list.append(end_page_item)

def remove_hanja(text):
    hanja_pattern = re.compile(r'[\u4E00-\u9FFF\u3400-\u4DBF]')
    cleaned_text = hanja_pattern.sub('', text)
    return cleaned_text


def split_number_string(s):
    #match = re.match(r"((?:\d+\.)+)(\S+)", s)'
    match = re.match(r"((?:\d+\.(?!\d))+)\s*(\S+)", s)
    if match:
        number_part = match.group(1)
        string_part = match.group(2)
        result = number_part + " "+string_part 
        result = re.sub(r'。', ' ', result)
        result = re.sub(r'\s{2,}', ' ', result)
        return result
    else:
        return ""
    

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pages.append(page.get_text())
    new_pages = []
    for page in pages:
    
        page = remove_hanja(page)
        page = page.replace("\t", " ").replace("\n", " ").strip().lower()
        #page = page.replace("-", " ")
        page = re.sub(r'。', ' ', page)
        page = re.sub(r'\s{2,}', ' ', page).strip()

        new_pages.append(page)

    return new_pages

def split_paragraph_pages(pages, section):
    result = []
    current_section = 0
    txt = ""   
    past_words = ""
    for page in pages:
        while current_section <= len(section) - 1:
            preprocess_section = remove_hanja(section[current_section])
            preprocess_section = preprocess_section.replace("\t", " ").replace("\n"," ").strip().lower()
            #preprocess_section = preprocess_section.replace("-", " ")
            preprocess_section = preprocess_section.replace('。', " ")
            section_list = re.sub(r'\s{2,}', ' ', preprocess_section).strip().split(' ')
            section_list = [section_list]
                        
            section2 = split_number_string(section[current_section])
            if (len(section2)  != 0):
                preprocess_section2 = remove_hanja(section2)
                preprocess_section2 = preprocess_section2.replace("\t", " ").replace("\n"," ").strip().lower()
                preprocess_section2 = preprocess_section2.replace("-", " ")
                section_list2 = re.sub(r'\s{2,}', ' ', preprocess_section2).strip().split(' ')
                section_list.append(section_list2)
            # Find the position of the section in the current page
            words = page.lower().split()
            start_idx = 0
            found = False
            break_point = True

            for s in section_list:
                for i in range(len(words)):
                    #print(words[i:i+len(s)], "와 ", s, "비교중")
                    if words[i:i+len(s)] == s:
                        found = True
                        start_idx = i
                        break

            if found:
                # If section is found in the middle of the page
                part1 = ' '.join(words[:start_idx])
                part2 = ' '.join(words[start_idx:])

                if (part1.strip() or txt.strip()):
                    result.append(txt.strip()+part1.strip())
                
                # Set the remaining part of the page to be processed
                page = part2
                txt = ""
                current_section += 1
                
            else:
                break
        
        txt += page + "\n"

    if txt.strip():
        result.append(txt.strip())

    return result



def extract_text_from_docx(file_path):
    doc = Document(file_path)
    new_pages = []

    for element in doc.element.body:
        if element.tag == qn('w:p'):  # Paragraph
            paragraph = element
            text = paragraph.text
            text = remove_hanja(text)
            text = text.replace("\t", " ").replace("\n", " ").strip().lower()
            text = re.sub(r'\s{2,}', ' ', text).strip()
            new_pages.append(text)
        elif element.tag == qn('w:tbl'):  # Table
            table_text = ""
            table = element
            for row in table.findall('.//w:tr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                row_text = ""
                for cell in row.findall('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    cell_text = ''.join(node.text for node in cell if node.text)
                    row_text += cell_text + '\t'
                row_text = row_text.strip() + '\n'
                table_text += row_text
            table_text += '\n'
            new_pages.append(table_text)
    return new_pages

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    text = '\n'.join(full_text)

    result = text.split("\n")
    with open("total_text.txt", 'w', encoding='utf-8') as f:
        f.write(text)

    with open('temp_result.txt', 'w', encoding='utf-8') as outfile:
        outfile.write(text)
    return result


read_optimizaton_data('optimization_data')

tests = [
    "Conducted Emission - Voltage Method", #0
    "Conducted Emission - Current Method", #1
    "Radiated Emission", #2
    "Magnetic Field Emission", #3
    "CV test", #4
    "Transient Emission", #5
    "Bulk Current Injection: Immunity", #6
    "ALSE: Immunity", #7
    "ALSE: Radar Pulse", #8
    "Portable Transmitter Test", #9
    "Magnetic Field Immunity", #10
    "Transient Immunity: Supply lines", #11
    "Transient Immunity; Signal lines", #12
    "Electrostatic Discharge (ESD) - Unpowered Discharge", #13
    "Electrostatic Discharge (ESD) - Powered-up Discharge", #14
    "Test to be performed as shwon" #15
]

text_contents = [
    "", #0
    "", #1
    "", #2
    "", #3
    "", #4 
    "", #5
    "", #6
    "", #7
    "", #8
    "", #9 
    "", #10
    "", #11
    "", #12
    "", #13
    "", #14
    "" #15
]
for i in range(len(file_path_list)):
    file_path = file_path_list[i]
    file_type = file_type_list[i]
    start_page = start_page_list[i]
    end_page = end_page_list[i]
    sections = sections_list[i]
    answer = answer_list[i] 
    
    if(file_type == 'pdf'):
        pages = extract_text_from_pdf(file_path)
    elif(file_type == 'docx'):
        pages = extract_text_from_docx(file_path)
        start_page = 2 
        end_page = len(pages) 
    else :
        print("요청한 파일 형식이 잘못 되었습니다.")
        sys.exit(0)
    
    if(file_type == 'pdf'):
        result = split_paragraph_pages(pages[start_page-1:end_page], sections) 
    elif(file_type == 'docx'):
        result = split_paragraph_pages(pages, sections)
    result.pop(0)
    for idx in range(len(result)):
        for j in range(len(answer[idx])):
            text_contents[answer[idx][j]] += result[idx]

def get_top_keywords(text, top_n=50):
    stop_words = set(stopwords.words('english'))
    stop_words.add("kevin.huang2")
    stop_words.add("test")
    stop_words.add("")
    punctuation = set(string.punctuation)
    
    words = nltk.word_tokenize(text.lower())
    filtered_words = [word for word in words if word not in stop_words and word not in punctuation]
    
    word_counts = Counter(filtered_words)

    top_keywords = word_counts.most_common(top_n)

    return top_keywords

result_txt = ""

for i in range(len(text_contents)):
    result_txt += "----------------------------------------------\n"
    result_txt += f"<<  {tests[i]}  >>\n"
    result_txt += str(get_top_keywords(text_contents[i], 40))
    result_txt += "\n"
  
    print("----------------------------------------------")
    print(f"<<  {tests[i]}  >>")
    print(get_top_keywords(text_contents[i], 30))

with open("recommend_parameters.txt", 'w', encoding='utf-8') as f:
    f.write(result_txt)

    

