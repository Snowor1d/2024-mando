from bayes_opt import BayesianOptimization
import numpy as np
import copy
import os

import fitz 
import copy
import re
from docx import Document
import sys 
from docx.oxml.ns import qn
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from rake_nltk import Rake
import nltk

file_path_list = []
file_type_list = []
start_page_list = []
end_page_list = []
sections_list = []
answer_list = []

tests = [
    "Conducted Emission - Voltage Method", #0
    "Conducted Emission - Current Method", #1
    "Radiated Emission", #2
    "Magnetic Field Emission", #3
    "CV test", #4
    "Transient Emission", #5
    "Bulk Current Injection: Immunity", #6
    "ALSE: Immunity", #7
    "ALSE: Radar Pulse", #8
    "Portable Transmitter Test", #9
    "Magnetic Field Immunity", #10
    "Transient Immunity: Supply lines", #11
    "Transient Immunity; Signal lines", #12
    "Electrostatic Discharge (ESD) - Unpowered Discharge", #13
    "Electrostatic Discharge (ESD) - Powered-up Discharge", #14
    "Test to be performed as shwon" #15
]



def read_test_info_from_file(file_path):
    with open(file_path, 'r') as file:
        lines = file.readlines()

    # Parse each line, splitting by comma and stripping whitespace
    additional_test_info = [line.strip().split(', ') for line in lines]

    # Convert empty strings to empty lists
    additional_test_info = [info if info != [''] else [] for info in additional_test_info]

    return additional_test_info

# Usage


file_path = 'parameter_candidates.txt'
parameter_candidates = read_test_info_from_file(file_path)

def read_optimizaton_data(folder_path):
    for filename in os.listdir(folder_path):
        if(filename.endswith(".txt")):
            file_path = os.path.join(folder_path, filename)
            print(f"Reading data from {file_path}")
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()

                file_path_item = lines[0].strip()
                file_type_item = lines[1].strip()
                start_page_item = int(lines[2].strip())
                end_page_item = int(lines[3].strip())
                answers = []
                sections = []
                for i in range(4, len(lines), 2):
                    section_item = lines[i].strip()
                    if i+1<len(lines):    
                        answer_item = [int(x) for x in lines[i+1].strip().split()]
                    else:
                        answer_item = []
                    sections.append(section_item)
                    answers.append(answer_item)
                sections_list.append(sections)
                answer_list.append(answers)
                file_path_list.append(file_path_item)
                file_type_list.append(file_type_item)
                start_page_list.append(start_page_item)
                end_page_list.append(end_page_item)
    


def remove_hanja(text):
    hanja_pattern = re.compile(r'[\u4E00-\u9FFF\u3400-\u4DBF]')
    cleaned_text = hanja_pattern.sub('', text)
    return cleaned_text


def split_number_string(s):
    #match = re.match(r"((?:\d+\.)+)(\S+)", s)'
    match = re.match(r"((?:\d+\.(?!\d))+)\s*(\S+)", s)
    if match:
        number_part = match.group(1)
        string_part = match.group(2)
        result = number_part + " "+string_part 
        result = re.sub(r'。', ' ', result)
        result = re.sub(r'\s{2,}', ' ', result)
        return result
    else:
        return ""


def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pages.append(page.get_text())
    new_pages = []
    for page in pages:
    
        page = remove_hanja(page)
        page = page.replace("\t", " ").replace("\n", " ").strip().lower()
        #page = page.replace("-", " ")
        page = re.sub(r'。', ' ', page)
        page = re.sub(r'\s{2,}', ' ', page).strip()

        new_pages.append(page)

    return new_pages

def split_paragraph_pages(pages, section):
    result = []
    current_section = 0
    txt = ""   
    past_words = ""
    for page in pages:
        while current_section <= len(section) - 1:
            preprocess_section = remove_hanja(section[current_section])
            preprocess_section = preprocess_section.replace("\t", " ").replace("\n"," ").strip().lower()
            #preprocess_section = preprocess_section.replace("-", " ")
            preprocess_section = preprocess_section.replace('。', " ")
            section_list = re.sub(r'\s{2,}', ' ', preprocess_section).strip().split(' ')
            section_list = [section_list]
                        
            section2 = split_number_string(section[current_section])
            if (len(section2)  != 0):
                preprocess_section2 = remove_hanja(section2)
                preprocess_section2 = preprocess_section2.replace("\t", " ").replace("\n"," ").strip().lower()
                section_list2 = re.sub(r'\s{2,}', ' ', preprocess_section2).strip().split(' ')
                section_list.append(section_list2)
            # Find the position of the section in the current page
            words = page.lower().split()
            start_idx = 0
            found = False
            break_point = True

            for s in section_list:
                for i in range(len(words)):
                    #print(words[i:i+len(s)], "와 ", s, "비교중")
                    if words[i:i+len(s)] == s:
                        found = True
                        start_idx = i
                        break

            if found:
                # If section is found in the middle of the page
                part1 = ' '.join(words[:start_idx])
                part2 = ' '.join(words[start_idx:])

                if (part1.strip() or txt.strip()):
                    result.append(txt.strip()+part1.strip())
                
                # Set the remaining part of the page to be processed
                page = part2
                txt = ""
                current_section += 1
                
            else:
                break
        
        txt += page + "\n"

    if txt.strip():
        result.append(txt.strip())

    return result

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    new_pages = []

    for element in doc.element.body:
        if element.tag == qn('w:p'):  # Paragraph
            paragraph = element
            text = paragraph.text
            text = remove_hanja(text)
            text = text.replace("\t", " ").replace("\n", " ").strip().lower()
            text = re.sub(r'\s{2,}', ' ', text).strip()
            new_pages.append(text)
        elif element.tag == qn('w:tbl'):  # Table
            table_text = ""
            table = element
            for row in table.findall('.//w:tr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                row_text = ""
                for cell in row.findall('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    cell_text = ''.join(node.text for node in cell if node.text)
                    row_text += cell_text + '\t'
                row_text = row_text.strip() + '\n'
                table_text += row_text
            table_text += '\n'
            new_pages.append(table_text)
    return new_pages

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    text = '\n'.join(full_text)

    result = text.split("\n")
    with open("total_text.txt", 'w', encoding='utf-8') as f:
        f.write(text)

    with open('temp_result.txt', 'w', encoding='utf-8') as outfile:
        outfile.write(text)
    return result


def check_answer(file_path, file_type, start_page, end_page, sections, tests, answer, test_input):
    global average_correct_rate
    result_txt = ""
    if(file_type == 'pdf'):
        pages = extract_text_from_pdf(file_path)
    elif(file_type == 'docx'):
        pages = extract_text_from_docx(file_path)
        start_page = 2
        end_page = len(pages)
    else :
        print("요청한 파일 형식이 잘못 되었습니다.")
        sys.exit(0)
    if(file_type == 'pdf'):
        result = split_paragraph_pages(pages[start_page-1:end_page], sections)
    elif(file_type == 'docx'):
        result = split_paragraph_pages(pages, sections)
    result.pop(0)
    if(len(sections) == len(result)):
        print("주어진 section을 모두 찾았습니다.")
    else :
        if (len(sections[len(result)]) > 1):
            print(f"{sections[len(result)]}부터 찾지 못했습니다.")
        else:
            print("주어진 section을 모두 찾았습니다.")
    # for i in result:
    #     #print(len(i))
    #     #print("------------")

    stop_words = nltk.corpus.stopwords.words('english')
    add_stopwords = ['test', 'testing','requirements', 'shall']
    for i in add_stopwords:
        stop_words.append(i)
    # 키워드 추출 (RAKE)
    rake = Rake(stopwords = stop_words)
    section_keywords = []
    for part in result:
        rake.extract_keywords_from_text(part)
        keywords = rake.get_ranked_phrases()
        #print("keyword 개수 : ", len(keywords))
        section_keywords.append(' '.join(keywords))

    # TF-IDF 벡터화
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(section_keywords + test_input)

    # 유사도 계산
    cosine_sim = cosine_similarity(tfidf_matrix[:len(result)], tfidf_matrix[len(result):])
    #print(cosine_sim)
    # 각 섹션을 가장 유사한 테스트에 매칭
    matched_tests = []
    question_num = 0
    answer_num = 0
    wrong_text = ""
    for idx, similarities in enumerate(cosine_sim):
        if idx == len(result):
            break
        most_similar_test_idx = similarities.argmax()
        matched_tests.append((result[idx], tests[most_similar_test_idx]))
        if len(answer[idx])!=0:
            question_num += 1
            if( most_similar_test_idx in answer[idx]):
                print(f"{sections[idx]} : {tests[most_similar_test_idx]} : 정답")
                answer_num += 1
            else :
                wrong_text += "{:>60}".format(f"{sections[idx]}") 
                wrong_text += f": {tests[most_similar_test_idx]} : 오답\n"
                print(f"{sections[idx]} : {tests[most_similar_test_idx]} : 오답")
    print(f"정답률 : {(answer_num/question_num)*100}")
    result_txt += "{:>60}".format(f'{str(file_path)} 정답률 : {(answer_num/question_num)*100}\n')
    result_txt += "{:>60}".format(wrong_text) 
    result_txt += "{:>60}".format("-------------------------------------------\n")

    # 매칭 결과 출력
    i = 0
    # for section, test in matched_tests:
    #     print(f"{sections[i]} : {test}\n")
    #     i += 1
    #     if(i == len(sections)):
    #         break

    return (answer_num/question_num)*100

def objective(**kwargs):
    naive_selected = [
        "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", ""
    ]

    mat = []
    for i in range(len(parameter_candidates)):
        mat.append([kwargs.get(f't{i}_{j}', 0) for j in range(len(parameter_candidates[i]))])

    for i in range(len(mat)):
        naive_selected[i] = tests[i]
        for j in range(len(mat[i])):
            if (mat[i][j]>0.5):
                naive_selected[i] += (" "+ parameter_candidates[i][j].lower())
    print(naive_selected) 

    tests_input = copy.deepcopy(naive_selected)
    total_score = 0
    
    for i in range(len(file_path_list)):
        file_path = file_path_list[i]
        file_type = file_type_list[i]
        start_page = start_page_list[i]
        end_page = end_page_list[i]
        sections = sections_list[i]
        answer = answer_list[i]
        total_score += check_answer(file_path, file_type, start_page, end_page, sections, tests, answer, tests_input)
    
    return total_score/len(file_path_list)

read_optimizaton_data('optimization_data')
pbounds = {
    f"t{i}_{j}": (0, 1) for i, test_params in enumerate(parameter_candidates) for j, param in enumerate(test_params)
}

optimizer = BayesianOptimization(
    f = objective,
    pbounds = pbounds,
    verbose = 2,
    random_state = 1,
)

optimizer.maximize(
    init_points = 10,
    n_iter = 5,
)

best_params = optimizer.max['params']

selected_parameters = []

for i, test_params in enumerate(parameter_candidates):
    if not test_params:
        continue

    selected_test_params = [
        param for j, param in enumerate(test_params)
        if best_params[f"t{i}_{j}"] > 0.5
    ]
    
    if selected_test_params:
        selected_parameters.append(selected_test_params)

def save_to_txt(file_path, selected_parameters):
    with open(file_path, 'w') as file:
        for params in selected_parameters:
            line = ', '.join(params)
            file.write(line + "\n")

file_path = 'selected_parameters.txt'
save_to_txt(file_path, selected_parameters)

print(f"Selected parameters have been saved to {file_path}")